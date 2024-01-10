


import pandas as pd
import numpy as np
from datetime import date
import csv
import docx
import urllib.request
import time
import datetime
from alpaca.trading.client import TradingClient

alpaca_keys = pd.read_csv('alpaca_api_information.csv')

# keys required for stock historical data client
#paper_tf = pd.read_csv('paper_true_false.csv')

trading_client = TradingClient(alpaca_keys['API KEY'][0], alpaca_keys['API SECRET KEY'][0], paper=alpaca_keys['Portfolio Type'][0])



print("SUMMARY REPORT COMMENCING")

while True:

    try:
        
        urllib.request.urlopen('http://google.com') #Python 3.x
        #print("Internet Connection Detected - Sumamry Report")
        #time.sleep(1)
    
        def summary_report():
        
            #today = date.today()
            #print("Today is: ", today)
            
            today_initial_data = date.today()
            y = today_initial_data.strftime("%m-%d-%Y")
            NY = 'America/los_angeles'
            end=pd.Timestamp(y, tz=NY)   
            today = y
            #print("Record Date:", today)
            
            #print('--------------------')
            
            account = trading_client.get_account()
            last_equity = account.last_equity
            #print("Last Equity:", last_equity)
            #print('--------------------')
            current_equity = account.equity
            #print("Current Equity:", current_equity)
            
            #print('--------------------')
            
            
            compiled_activity_results = pd.read_csv('one_day_1_percentage_movements.csv')
            compiled_activity_results.pop('Unnamed: 0')
            ticker_append = []
            dollar_sum_append = []
            for items_tickers in compiled_activity_results['Ticker']:
                #print(items_tickers)
                ticker_append.append(items_tickers)
                current_ticker_loop_list = compiled_activity_results.drop(compiled_activity_results[compiled_activity_results.Ticker != items_tickers].index)
                dollar_sum = current_ticker_loop_list['Dollar']
                dollar_sum = sum(dollar_sum)
                #print(dollar_sum)
                dollar_sum_append.append(dollar_sum)
            ticker_df = pd.DataFrame(ticker_append)
            ticker_df.columns = ['Ticker']
            #print(ticker_df)
            dollar_df = pd.DataFrame(dollar_sum_append)
            dollar_df.columns = ['Total Gain/Loss']
            #print(dollar_df)
            combine = pd.concat([ticker_df,dollar_df], axis = 1)
            combine_post_dropped = combine.drop_duplicates()
            #print("Simmulated Individual Stock Gains:")
            #print(combine_post_dropped)
            #print('--------------------')
            combine_post_dropped['Total Gain/Loss'].to_numpy().tolist()
            
            
            simmulated_portfolio_gain = sum(combine_post_dropped['Total Gain/Loss'])
            #print('Simmulated Portfolio Gain ($):', simmulated_portfolio_gain)
            sim_per_gain = (simmulated_portfolio_gain/float(last_equity))*100
            #print('Simmulated Portfolio Gain (%):', sim_per_gain)
            #print('--------------------')
            
            portfolio_gain_loss = float(current_equity) - float(last_equity)
            
            act_per_gain = (float(portfolio_gain_loss)/float(last_equity))*100
            
            
            
            
            spread_cost = abs((simmulated_portfolio_gain - portfolio_gain_loss))
            #print('Spread Cost ($):', spread_cost)
            perc_spread_cost = (spread_cost/float(last_equity))*100
            #print('Spread Cost (%):', perc_spread_cost)
            
            #print('--------------------')
            
            spread_cost_per_stock = spread_cost/len(combine_post_dropped['Ticker'])
            #print('Average Spread Cost Per Stock ($):', spread_cost_per_stock)
            perc_spread_per_stock = (spread_cost_per_stock/float(last_equity))*100
            #print('Average Spread Cost Per Stock (%):', perc_spread_per_stock)
            
            #print('--------------------')
            
            actual = combine_post_dropped['Total Gain/Loss'] - spread_cost_per_stock
            
            actual_concat = pd.concat([combine_post_dropped['Ticker'], actual], axis = 1)
            #print("Actual Individual Stock Gains:")
            #print(actual_concat)
            
            #print('--------------------')
            
            #print('Actual Portfolio Gain ($):', portfolio_gain_loss)
            #print('Actual Portfolio Gain (%):', act_per_gain)
            
            
            #print('--------------------')
            
            #print('Number of Stocks Used in Trading:', len(combine_post_dropped['Ticker']))
            ticker_parameter_grades  = pd.read_csv('ticker_parameter_grades.csv')
            ticker_list = ticker_parameter_grades['Ticker'].to_numpy().tolist()
            number_of_stocks_initial = len(ticker_list)
            #print("Number of Stocks Chosen Before Trading:", number_of_stocks_initial)
            #print("Stocks Chosen Before Trading:", ticker_list)
            #print("Stocks Used In Trading:", combine_post_dropped['Ticker'].to_numpy().tolist())
            #print('--------------------')
            
            
            
            
            mydoc = docx.Document()
            
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Today is: "+str(today))
            
            mydoc.add_paragraph("Last Equity: "+str(last_equity))
            mydoc.add_paragraph("Current Equity: "+str(current_equity))
            
            table = mydoc.add_table(rows=1, cols=2)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Ticker'
            row[1].text = 'Total Gain/Loss'
            
            for word_ticker, word_numbers in zip(combine_post_dropped['Ticker'], combine_post_dropped['Total Gain/Loss']):
              
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = word_ticker
                row[1].text = str(word_numbers)
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Simmulated Portfolio Gain ($): "+str(simmulated_portfolio_gain))
            mydoc.add_paragraph("Simmulated Portfolio Gain (%): "+str(sim_per_gain))
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Spread Cost ($): "+str(spread_cost))
            mydoc.add_paragraph("Spread Cost (%): "+str(perc_spread_cost))
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Average Spread Cost Per Stock ($): "+str(spread_cost_per_stock))
            mydoc.add_paragraph("Average Spread Cost Per Stock (%): "+str(perc_spread_per_stock))
            #mydoc.add_paragraph("--------------------")
            table = mydoc.add_table(rows=1, cols=2)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Ticker'
            row[1].text = 'Total Gain/Loss'
            
            for word_ticker, word_numbers in zip(actual_concat['Ticker'], actual_concat['Total Gain/Loss']):
              
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = word_ticker
                row[1].text = str(word_numbers)
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Actual Portfolio Gain ($): "+str(portfolio_gain_loss))
            mydoc.add_paragraph("Actual Portfolio Gain (%): "+str(act_per_gain))
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Number of Stocks Used in Trading: "+str(len(combine_post_dropped['Ticker'])))
            mydoc.add_paragraph("Number of Stocks Chosen Before Trading: "+str(number_of_stocks_initial))
            #mydoc.add_paragraph("--------------------")
            mydoc.add_paragraph("Stocks Chosen Before Trading: "+str(ticker_list))
            mydoc.add_paragraph("Stocks Used In Trading: "+str(combine_post_dropped['Ticker'].to_numpy().tolist()))
            #mydoc.add_paragraph("--------------------")
            
            
            mydoc.save("summary_report_"+str(today)+".docx")
            
        print("SUMMARY REPORT COMPLETE")
                
        summary_report()  
        
        break
        
    except:
        print("No Internet Connection, Attempting Reconnection... - Sumamry Report")
        time.sleep(0.5)
  




